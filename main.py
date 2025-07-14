import os
import re
import requests
import datetime
from dotenv import load_dotenv
from requests.auth import HTTPBasicAuth
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from github import Github
from jira import JIRA
from openai import OpenAI
from typing import List, Dict

# === CARGAR VARIABLES DE ENTORNO ===
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key)

GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
GITHUB_REPO = os.getenv("GITHUB_REPO")
JIRA_EMAIL = os.getenv("JIRA_EMAIL")
JIRA_TOKEN = os.getenv("JIRA_TOKEN")
JIRA_URL = os.getenv("JIRA_URL")
JIRA_PROJECT_KEY = os.getenv("JIRA_PROJECT_KEY")
JIRA_ISSUE_TYPE = os.getenv("JIRA_ISSUE_TYPE")


def fetch_release_notes(tag):
    g = Github(GITHUB_TOKEN)
    repo = g.get_repo(GITHUB_REPO)
    try:
        release = repo.get_release(tag)
        global release_date
        release_date = release.published_at.strftime("%Y-%m-%d") if release.published_at else datetime.datetime.today().strftime("%Y-%m-%d")

        if release.body and release.body.strip():
            return release.body

        print("⚠️ No hay release notes. Buscando commits relacionados al tag...")
        tag_obj = repo.get_git_ref(f"tags/{tag}")
        tag_sha = tag_obj.object.sha
        tag_commit = repo.get_commit(tag_sha)
        commits = repo.get_commits(since=tag_commit.commit.author.date)

        release_notes = ""
        for commit in commits:
            if tag in commit.commit.message:
                release_notes += f"- {commit.commit.message}\n"

        return release_notes or f"No release notes or matching commits found for tag {tag}."

    except Exception as e:
        print("❌ Error al obtener el release de GitHub:", str(e))
        return None


def extract_sections(content):
    bug_fixes, known_issues, other_lines = [], [], []
    current_section = None
    for line in content.splitlines():
        stripped = line.strip().lower()
        if "bug fixes" in stripped:
            current_section = "bug"
        elif "known issues" in stripped:
            current_section = "issues"
        elif not stripped:
            continue
        else:
            if current_section == "bug":
                bug_fixes.append(line)
            elif current_section == "issues":
                known_issues.append(line)
            else:
                other_lines.append(line)
    return bug_fixes, known_issues, other_lines


def fetch_jira_ticket_details(ticket_id):
    url = f"{JIRA_URL}/rest/api/3/issue/{ticket_id}"
    auth = HTTPBasicAuth(JIRA_EMAIL, JIRA_TOKEN)
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers, auth=auth)
    if response.status_code == 200:
        data = response.json()
        fields = data["fields"]
        return {
            "id": ticket_id,
            "summary": fields.get("summary", ""),
            "status": fields.get("status", {}).get("name", ""),
            "url": f"{JIRA_URL}/browse/{ticket_id}",
            "epic": fields.get("customfield_10014", "Other") or "Other"
        }
    return {
        "id": ticket_id,
        "summary": "No encontrado",
        "status": "Desconocido",
        "url": f"{JIRA_URL}/browse/{ticket_id}",
        "epic": "Other"
    }

def build_rich_adf_description(release_body, tickets, summary_text=None):
    import re

    def make_link(text, url):
        return {
            "type": "text",
            "text": text,
            "marks": [{"type": "link", "attrs": {"href": url}}]
        }

    content = []

    # Header
    content.append({
        "type": "heading",
        "attrs": {"level": 2},
        "content": [{"type": "text", "text": "📦 Release Notes"}]
    })

    # Optional compare URL
    first_line = release_body.splitlines()[0] if release_body else ""
    match = re.search(r"\[(.*?)\]\((https:\/\/github\.com\/[^)]+)\)\s*\((\d{4}-\d{2}-\d{2})\)", first_line)
    if match:
        _, compare_url, date = match.groups()
        content.append({
            "type": "paragraph",
            "content": [
                {"type": "text", "text": "🔍 "},
                make_link("View Code Changes", compare_url),
                {"type": "text", "text": f" – {date}"}
            ]
        })

    # Extract sections
    lines = release_body.splitlines()
    features, bugs = [], []
    current = None

    for line in lines:
        line = line.strip()
        if not line or line == first_line:
            continue
        if line.lower().startswith("### features"):
            current = "features"
            continue
        elif line.lower().startswith("### bug fixes"):
            current = "bugs"
            continue
        if current == "features":
            features.append(line)
        elif current == "bugs":
            bugs.append(line)

    def render_section(title, lines):
        content.append({
            "type": "heading",
            "attrs": {"level": 3},
            "content": [{"type": "text", "text": title}]
        })

        for line in lines:
            paragraph = {"type": "paragraph", "content": []}

            ticket_matches = re.findall(r"\[(WEBTV-\d+|CWB-\d+)\]", line)
            commit_links = re.findall(r"\((https:\/\/github\.com\/[^\s]+\/commit\/[a-f0-9]{7,40})\)", line)
            pr_links = re.findall(r"\((https:\/\/github\.com\/[^\s]+\/issues\/\d+)\)", line)

            # Clean visible text
            clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            clean_text = re.sub(r"\[(WEBTV-\d+|CWB-\d+)\]", "", clean_text)
            clean_text = re.sub(r"\(https?:\/\/[^\s)]+\)", "", clean_text)
            clean_text = re.sub(r"https?:\/\/[^\s]+", "", clean_text)
            clean_text = re.sub(r"\(\s*\[+\s*\]+\s*\)", "", clean_text)
            clean_text = re.sub(r"[\[\]()]", "", clean_text)
            clean_text = re.sub(r"^[•\-\.\s:—]+", "", clean_text).strip()

            if not clean_text:
                continue

            if not clean_text.endswith("."):
                clean_text += "."

            paragraph["content"].append({"type": "text", "text": f"• {clean_text} "})

            # Add ticket links
            if ticket_matches:
                for t in ticket_matches:
                    paragraph["content"].append(make_link(f"🔗 {t}", f"https://televisaunivision.atlassian.net/browse/{t}"))

            # Add PR/commit links
            if pr_links:
                paragraph["content"].append({"type": "text", "text": " – "})
                paragraph["content"].append(make_link("[PR]", pr_links[0]))
            if commit_links:
                paragraph["content"].append({"type": "text", "text": " "})
                paragraph["content"].append(make_link("[Code]", commit_links[0]))

            content.append(paragraph)

            # Optional blockCard
            for t in ticket_matches:
                content.append({
                    "type": "blockCard",
                    "attrs": {
                        "url": f"https://televisaunivision.atlassian.net/browse/{t}"
                    }
                })

    if features:
        render_section("🚀 Features", features)
    if bugs:
        render_section("🐞 Bug Fixes", bugs)

    # 📌 Summary
    summary_paragraphs = [
        {
            "type": "heading",
            "attrs": {"level": 3},
            "content": [{"type": "text", "text": "📌 Summary"}]
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": f"• {len(features)} features delivered 🚀"}]
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": f"• {len(bugs)} bugs resolved 🐞"}]
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "• 0 known issues ⚠️"}]
        }
    ]

    content.extend(summary_paragraphs)

    return {
        "type": "doc",
        "version": 1,
        "content": content
    }

def markdown_to_adf_paragraphs(md_text):
    paragraphs = []
    lines = md_text.strip().split("\n")
    jira_pattern = re.compile(r'\b([A-Z]+-\d+)\b')  # más flexible
    bold_pattern = re.compile(r"\*\*(.*?)\*\*")

    for line in lines:
        if not line.strip():
            continue

        content = []
        line = line.strip()

        # ✅ Limpieza de enlaces Markdown
        line = re.sub(r"\[\]\(https:\/\/[^\)]+\)", "", line)
        line = re.sub(r"\(\[\]\(https:\/\/[^\)]+\)\)", "", line)
        line = re.sub(r"\(https:\/\/[^\)]+\)", "", line)

        # ✅ Eliminar paréntesis vacíos
        line = re.sub(r"\(\s*\[?\s*\]?\s*\)", "", line)
        line = re.sub(r"\(\s*\)", "", line)
        line = line.replace("()", "").replace("[]", "").replace("  ", " ").strip()

        # ✅ Detectar línea que empieza con emoji + título (y ponerlo en negrita)
        emoji_heading = re.match(r"^([\U0001F300-\U0001FAFF\u2600-\u26FF\u2700-\u27BF]{1,2})\s+(.+)", line)
        if emoji_heading:
            emoji, text = emoji_heading.groups()
            content = [
                {"type": "text", "text": f"{emoji} ", "marks": []},
                {"type": "text", "text": text.strip(), "marks": [{"type": "strong"}]}
            ]
            paragraphs.append({"type": "paragraph", "content": content})
            continue

        # ✅ Títulos (encabezados)
        if line.startswith("### "):
            text = line[4:]
            content.append({"type": "text", "text": text, "marks": [{"type": "strong"}]})
            paragraphs.append({"type": "heading", "attrs": {"level": 3}, "content": content})
            continue
        elif line.startswith("## "):
            text = line[3:]
            content.append({"type": "text", "text": text, "marks": [{"type": "strong"}]})
            paragraphs.append({"type": "heading", "attrs": {"level": 2}, "content": content})
            continue
        elif line.startswith("# "):
            text = line[2:]
            content.append({"type": "text", "text": text, "marks": [{"type": "strong"}]})
            paragraphs.append({"type": "heading", "attrs": {"level": 1}, "content": content})
            continue

        # ✅ Extraer ticket ID
        ticket_ids = jira_pattern.findall(line)

        # ✅ Separar negritas del resto
        parts = bold_pattern.split(line)
        for i, part in enumerate(parts):
            if i % 2 == 1:
                # Es texto en negrita
                content.append({"type": "text", "text": "• " + part.strip() + ": ", "marks": [{"type": "strong"}]})
            else:
                if part.strip():
                    content.append({"type": "text", "text": part.strip()})

        # ✅ Agregar párrafo con texto principal
        paragraphs.append({"type": "paragraph", "content": content})

        # ✅ Agregar solo un botón de ticket, si hay ID válido
        if ticket_ids:
            ticket_url = f"https://televisaunivision.atlassian.net/browse/{ticket_ids[0]}"
            paragraphs.append({
                "type": "paragraph",
                "content": [{
                    "type": "text",
                    "text": "🔗 View Ticket",
                    "marks": [{"type": "link", "attrs": {"href": ticket_url}}]
                }]
            })

    return paragraphs


def parse_github_release_notes(release_notes):
    """
    Transforma el release.body plano de GitHub en una lista estructurada
    para usarla en generate_friendly_summary y otras funciones.
    """
    import re

    structured = []
    lines = release_notes.splitlines()

    current_section = None  # puede ser "feature" o "bug"

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detectar sección (Features o Bug Fixes)
        if "features" in line.lower():
            current_section = "feature"
            continue
        elif "bug fixes" in line.lower():
            current_section = "bug"
            continue

        # Saltar líneas de encabezado
        if line.startswith("#"):
            continue

        # Extrae el ticket si hay
        match = re.search(r'\b([A-Z]+-\d+)\b', line)
        ticket_id = match.group(1) if match else "UNKNOWN"

        # Si no hay sección detectada, hacer heurística
        l = line.lower()
        if current_section:
            item_type = current_section
        elif any(w in l for w in ['fix', 'bug', 'resolve']):
            item_type = 'bug'
        elif any(w in l for w in ['add', 'implement', 'migrate', 'move', 'create', 'remove gate', 'rsc']):
            item_type = 'feature'
        else:
            item_type = 'feature'

        structured.append({
            "type": item_type,
            "title": line[:80],
            "description": line,
            "ticket_id": ticket_id
        })

    return structured


def create_jira_issue(summary, adf_description, tickets, release_notes):
    url = f"{JIRA_URL}/rest/api/3/issue"
    headers = {"Accept": "application/json", "Content-Type": "application/json"}
    auth = HTTPBasicAuth(JIRA_EMAIL, JIRA_TOKEN)

    # ✅ Construir structured_release_notes correctamente
    structured_release_notes = []

    if not tickets or len(tickets) < 5:
        structured_release_notes = parse_github_release_notes(release_notes)
    else:
        for ticket in tickets:
            raw_summary = ticket.get("summary", "").strip()
            ticket_id = ticket.get("key", ticket.get("id", "")).strip()
            if not ticket_id:
                continue

            issue_type_raw = raw_summary.lower()

            # Palabras clave para detectar features
            feature_keywords = [
                "add", "rsc", "migrate", "move", "create", "implement",
                "profile", "modal", "pill", "tab", "ui", "gate", "flow",
                "state machine", "component", "split screen", "login", "logout", "subscribe"
            ]

            is_feature = any(word in issue_type_raw for word in feature_keywords)
            is_bug = any(word in issue_type_raw for word in ["fix", "bug", "error", "issue", "broken", "not working"])

            if is_feature:
                issue_type = "feature"
            elif is_bug:
                issue_type = "bug"
            else:
                issue_type = "feature"  # Default

            structured_release_notes.append({
                "type": issue_type,
                "title": raw_summary,
                "description": raw_summary,
                "ticket_id": ticket_id
            })

    # ✅ Generar resumen amigable con GPT-4
    version_tag = summary.replace("Release Document for ", "").strip()
    friendly_response = generate_friendly_summary(version_tag, structured_release_notes)

    # ✅ Convertir resumen a párrafos ADF
    intro_paragraphs = [
        {
            "type": "heading",
            "attrs": {"level": 2},
            "content": [{"type": "text", "text": "🧠 Summary Highlights"}]
        }
    ] + friendly_response + [
        {"type": "paragraph", "content": [{"type": "text", "text": ""}]}
    ]

    # ✅ Insertar el resumen al inicio
    adf_description["content"] = intro_paragraphs + adf_description["content"]

    # ✅ Crear el ticket en Jira
    payload = {
        "fields": {
            "project": {"key": JIRA_PROJECT_KEY},
            "summary": summary,
            "description": adf_description,
            "issuetype": {"name": JIRA_ISSUE_TYPE},
            "labels": ["release-documentation"],
        }
    }

    response = requests.post(url, headers=headers, auth=auth, json=payload)
    if response.status_code == 201:
        key = response.json().get("key")
        print("✅ Ticket creado:", key)
        return key
    else:
        print("❌ Error al crear ticket Jira:", response.status_code, response.text)
        return None

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_better_word(markdown_lines, output_path, tickets_info, version_tag, release_date, release_owner):
    doc = Document()
    title = doc.add_paragraph()
    title_run = title.add_run(f"📄 Release Management Document - {version_tag}")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    doc.add_paragraph(f"📦 Release Version: {version_tag}").bold = True
    doc.add_paragraph(f"📅 Release Date: {release_date}").bold = True
    doc.add_paragraph(f"👤 Release Owner: {release_owner}").bold = True
    doc.add_paragraph()

    summary_added = False
    detail_added = False
    seen_lines = set()

    for line in markdown_lines:
        line = line.strip()
        if line.lower() == "bug fixes":
            doc.add_paragraph("🐞 Bug Fixes", style="Heading 2")
            continue
        if line.lower() == "known issues":
            doc.add_paragraph("⚠️ Known Issues", style="Heading 2")
            continue

        if "[CWB-" in line or "[WEBTV-" in line:
            clean = line.replace("**", "").replace("*", "").strip()
            if clean in seen_lines:
                continue
            seen_lines.add(clean)

            if not summary_added:
                doc.add_paragraph("🔹 Summary of Changes", style="Heading 2")
                summary_added = True
            doc.add_paragraph(clean, style='List Bullet')

            if not detail_added:
                doc.add_paragraph("📋 Detailed Release Notes", style="Heading 2")
                detail_added = True

            p = doc.add_paragraph(style='List Number')
            found = False
            for t in tickets_info:
                if t["id"] in clean:
                    add_hyperlink(p, f"🔗 {clean}", t["url"])
                    found = True
                    break
            if not found:
                p.add_run(clean)

            ticket_data = next((t for t in tickets_info if t["id"] in clean), None)
            if ticket_data:
                p_status = doc.add_paragraph(f"Status: {ticket_data['status']}")
                p_status.paragraph_format.left_indent = Inches(0.25)

                if ticket_data.get("description"):
                    p_desc = doc.add_paragraph(f"Description: {ticket_data['description']}")
                    p_desc.paragraph_format.left_indent = Inches(0.25)

    doc.save(output_path)
    print(f"📄 Documento Word guardado: {output_path}")


def generate_release_doc_with_gpt(version_tag, release_notes, tickets_info):
    prompt = f"""
You're an expert release engineer. Based on the following GitHub release notes and Jira ticket details, generate a professional Release Management Document for version {version_tag}.

Include:
1. Summary of Changes (brief, bullet points)
2. Detailed Release Notes (with ticket key, description, and status)
3. Known Issues (if any)
4. Notes for QA (if needed)

GitHub Release Notes:
{release_notes}

Jira Ticket Details:
{tickets_info}
"""
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant who writes release documentation."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.4
    )
    return response.choices[0].message.content


def answer_question_with_gpt(question, release_notes, tickets_info):
    context = f"""
Release Notes:
{release_notes}

Jira Tickets:
{tickets_info}
"""
    prompt = f"""
Context:
{context}

User Question:
{question}

Answer the user's question based on the provided release and ticket data. Be concise and accurate.
"""
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that answers questions about software releases."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content

def generate_friendly_summary(version_tag, structured_release_notes):
    features = []
    bugs = []

    for item in structured_release_notes:
        title = item["title"].strip(" -•:—.")
        description = item["description"].strip(" -•:—.")
        ticket_id = item["ticket_id"]

        sentence = f"{title}: {description}" if description and description != title else title
        line = f"- {sentence} ({ticket_id})"
        if item["type"] == "feature":
            features.append(line)
        else:
            bugs.append(line)

    features_text = "\n".join(features[:4]) if features else "- No new features in this release."
    bugs_text = "\n".join(bugs[:4]) if bugs else "- No bug fixes in this release."

    prompt = f"""
You are a senior product release writer. Write a professional and user-friendly release summary for version {version_tag}.

Format:

🎉 What's New in Version {version_tag}

🚀 New Features & Improvements
{features_text}

🛠️ Bug Fixes
{bugs_text}

📌 Summary
Write a 3-4 sentence summary explaining the value of this release using the changes above.

Guidelines:
- Include short but clear explanations (1 line) per item.
- Ticket IDs must be shown in parentheses, e.g., (TVAPP-1234)
- NO empty (), no raw links, no Markdown formatting.
- Use proper punctuation and polish grammar.
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that writes clean Jira summaries."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )

    summary_text = response.choices[0].message.content.strip()

    # ✅ Convert the AI summary to proper ADF with bold titles
    return markdown_to_adf_paragraphs(summary_text)

# =============================
# SOLO SE EJECUTA SI CORRES main.py DIRECTAMENTE
# =============================
if __name__ == "__main__":
    version_tag = input("Escribe el tag de versión (ej: v1.109.0-beta.4): ").strip()
    if not version_tag:
        print("❌ No ingresaste ningún tag de versión.")
        exit(1)

    safe_version_tag = version_tag.replace("/", "-").replace("\\", "-")
    md_filename = f"release_{safe_version_tag}.md"
    word_filename = f"release_{safe_version_tag}.docx"

    release_notes = fetch_release_notes(version_tag)
    if not release_notes:
        print("❌ No se encontraron release notes.")
        exit(1)

    bug_fixes, known_issues, other_lines = extract_sections(release_notes)
    markdown_lines = other_lines[:]
    if bug_fixes:
        markdown_lines.append("Bug Fixes")
        markdown_lines.extend(bug_fixes)
    if known_issues:
        markdown_lines.append("Known Issues")
        markdown_lines.extend(known_issues)

    ticket_ids = list(set(re.findall(r'(CWB-\d+|WEBTV-\d+)', release_notes)))
    tickets_info = [fetch_jira_ticket_details(tid) for tid in ticket_ids]

    ai_generated_text = generate_release_doc_with_gpt(version_tag, release_notes, tickets_info)

    with open(md_filename, "w", encoding="utf-8") as f:
        f.write(ai_generated_text)
        print(f"📄 Documento AI guardado: {md_filename}")

    generate_better_word(markdown_lines, word_filename, tickets_info, version_tag, release_date, JIRA_EMAIL)


    friendly_summary = generate_friendly_summary(version_tag, structured_release_notes)
    create_jira_issue(f"Release Management Document - {version_tag}", friendly_summary, tickets_info)


    while True:
        user_question = input("\n¿Tienes alguna pregunta sobre este release? (Escribe 'no' para salir): ")
        if user_question.lower() in ['no', 'n']:
            print("👋 Saliendo del modo de preguntas.")
            break
        answer = answer_question_with_gpt(user_question, release_notes, tickets_info)
        print("\n🤖 Respuesta AI:")
        print(answer)
