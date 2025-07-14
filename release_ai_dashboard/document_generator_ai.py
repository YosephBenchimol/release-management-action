from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
from gpt_utils import add_hyperlink
import os
import json

load_dotenv()
client = OpenAI()

def translate_to_english_if_needed(text_list, label):
    """Detect and translate a list of items to English using GPT-4 if needed."""
    if not text_list:
        return []
    joined_text = "\n".join([line if isinstance(line, str) else "" for line in text_list])
    prompt = f"""
You are an assistant that ensures technical content is written in clear, professional English.

The following is a list of items labeled '{label}'. If they are in Spanish, translate them to English. If they are already in English, keep them as is. Return each item on a new line:

{joined_text}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a translation assistant."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1
    )
    return response.choices[0].message.content.strip().split("\n")

def smart_generate_sections(raw_release_notes: str) -> dict:
    prompt = f"""
You are an expert technical writer. Classify and rewrite the following release notes content into clean, structured sections:
- Summary of Changes (1 paragraph max)
- Detailed Release Notes (list format, include ticket ID if present)
- Known Issues (if none, say 'No known issues at this time.')

Release Input:
{raw_release_notes}

Output must be in this JSON format:
{{
  "summary": "...",
  "detailed_notes": ["...", "..."],
  "known_issues": ["..."]
}}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    try:
        return json.loads(response.choices[0].message.content)
    except Exception:
        return {"summary": "", "detailed_notes": [], "known_issues": []}

def generate_release_doc_with_gpt(release_notes, jira_info):
    prompt = f"""
You are a Release Management Assistant.

Based on the following release notes and Jira information, generate a professional and structured Release Management Document:

Release Notes:
{release_notes}

Jira Tickets Info:
{jira_info}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a Release Management Document Assistant."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content

def save_to_word(content, filename):
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    output_path = os.path.join("static", filename)
    doc.save(output_path)
    return output_path

def generate_summary_paragraph(summary_data: list[str]) -> str:
    paragraph = "📌 **Summary** — "

    for sentence in summary_data:
        sentence = sentence.strip().rstrip(".")
        if "profile" in sentence.lower():
            paragraph += f"🚀 {sentence.capitalize()}. "
        elif "login" in sentence.lower() or "authentication" in sentence.lower():
            paragraph += f"🔐 {sentence.capitalize()}. "
        elif "bug" in sentence.lower() or "issue" in sentence.lower():
            paragraph += f"🐞 {sentence.capitalize()}. "
        elif "cookie" in sentence.lower() or "security" in sentence.lower():
            paragraph += f"🍪 {sentence.capitalize()}. "
        else:
            paragraph += f"{sentence.capitalize()}. "

    return paragraph.strip()

def generate_structured_release_doc(filename, release_info):
    doc = Document()

    # AI Enhancement: classify and rewrite sections from release_body
    release_body = release_info.get('release_body', '')
    if release_body:
        ai_sections = smart_generate_sections(release_body)
        ai_summary = ai_sections.get("summary", "")
        ai_details = ai_sections.get("detailed_notes", [])
        ai_known = ai_sections.get("known_issues", [])
    else:
        ai_summary, ai_details, ai_known = "", [], []

    # Translate content to English
    summary = translate_to_english_if_needed(release_info.get('summary', []), "Summary of Changes")
    checklist = translate_to_english_if_needed(release_info.get('checklist', []), "Deployment Checklist")
    stakeholders = translate_to_english_if_needed(release_info.get('stakeholders', []), "Stakeholders & Approvals")

    # Title and metadata
    doc.add_heading('📄 Release Management Document', level=1)
    doc.add_paragraph(f"Project Name: {release_info.get('project_name', 'N/A')}")
    doc.add_paragraph(f"Version: {release_info.get('version', 'N/A')}")
    doc.add_paragraph(f"Release Date: {release_info.get('release_date', 'N/A')}")
    doc.add_paragraph(f"Prepared By: {release_info.get('prepared_by', 'N/A')}")
    doc.add_paragraph()

    # Section 1: Summary of Changes
    doc.add_heading("1. Summary of Changes", level=2)
    if ai_summary:
        doc.add_paragraph(ai_summary)
    elif summary:
        formatted_summary = generate_summary_paragraph(summary)
        doc.add_paragraph(formatted_summary, style="Normal")
    else:
        doc.add_paragraph("No summary available.")
    doc.add_paragraph()

    # Section 2: Detailed Release Notes
    doc.add_heading("2. Detailed Release Notes", level=2)
    details = ai_details or release_info.get('details', [])
    if details:
        if isinstance(details[0], dict):  # structured Jira-style notes
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Ticket'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Link'
            for detail in details:
                row_cells = table.add_row().cells
                row_cells[0].text = detail.get('ticket', '')
                row_cells[2].text = detail.get('status', '')

                # Write description + inline link in same cell
                desc = detail.get('description', '')
                link = detail.get('link', '')
                p = row_cells[1].paragraphs[0]
                p.add_run(desc + " ")
                add_hyperlink(p, "🔗 View Ticket", link)

                # Leave 4th cell empty
                row_cells[3].text = ''
        else:  # AI-style list notes
            for note in details:
                if "New Features & Improvements" in note:
                    p = doc.add_paragraph()
                    run = p.add_run("🚀 • New Features & Improvements —")
                    run.bold = True
                    run.font.size = Pt(13)
                elif "Bug Fixes" in note:
                    p = doc.add_paragraph()
                    run = p.add_run("🛠️ • Bug Fixes —")
                    run.bold = True
                    run.font.size = Pt(13)
                else:
                    doc.add_paragraph(f"- {note}", style="List Bullet")
    else:
        doc.add_paragraph("No detailed release notes provided.")
    doc.add_paragraph()

    # Section 3: Known Issues
    doc.add_heading("3. Known Issues", level=2)
    known_issues = ai_known or release_info.get('known_issues', [])
    if known_issues:
        for issue in known_issues:
            doc.add_paragraph(f"- {issue}", style="List Bullet")
    else:
        doc.add_paragraph("No known issues reported in this release.")
    doc.add_paragraph()

    # Section 4: Deployment Checklist
    doc.add_heading("4. Deployment Checklist", level=2)
    if checklist:
        for item in checklist:
            doc.add_paragraph(f"✅ {item}", style="List Bullet")
    else:
        doc.add_paragraph("No deployment checklist provided.")
    doc.add_paragraph()

    # Section 5: Stakeholders & Approvals
    doc.add_heading("5. Stakeholders & Approvals", level=2)
    if stakeholders:
        for stakeholder in stakeholders:
            doc.add_paragraph(f"- {stakeholder}", style="List Bullet")
    else:
        doc.add_paragraph("No stakeholder information provided.")
    doc.add_paragraph()

    # Section 6: Supporting Links
    doc.add_heading("6. Supporting Links", level=2)
    links = release_info.get('links', [])
    if links:
        for label, url in links:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ")
            add_hyperlink(p, url, url)
    else:
        doc.add_paragraph("Pending technical documentation link.")
    doc.add_paragraph()

    # Footer
    doc.add_paragraph("This document was automatically generated using GPT-4 AI Release Assistant.", style="Intense Quote")

    # Save document
    output_path = os.path.join("static", filename)
    doc.save(output_path)
    return output_path