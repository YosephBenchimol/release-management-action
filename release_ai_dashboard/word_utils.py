from docx import Document
from docx.shared import Pt

def save_word_from_text(text, filename):
    doc = Document()
    for line in text.splitlines():
        if line.strip().startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=1)
        elif line.strip().startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line.strip()).style.font.size = Pt(11)
    doc.save(filename)
