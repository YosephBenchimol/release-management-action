from dotenv import load_dotenv
import os
import re
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    print("❌ OPENAI_API_KEY no está disponible en el entorno.")
else:
    print("✅ OPENAI_API_KEY cargado correctamente.")

client = OpenAI(api_key=api_key)

def generate_release_doc_with_gpt(version_tag, release_notes, tickets_info):
    jira_base_url = "https://televisaunivision.atlassian.net/browse"

    enriched_tickets = []
    for line in tickets_info.strip().split("\n"):
        match = re.match(r"(CWB-\d+): (.*)", line)
        if match:
            ticket_id, description = match.groups()
            hyperlink_line = f"- [{ticket_id}]({jira_base_url}/{ticket_id}): {description}"
            enriched_tickets.append(hyperlink_line)
        else:
            enriched_tickets.append(line)

    linked_ticket_text = "\n".join(enriched_tickets)

    prompt = f"""
Act as a senior technical engineer. Write a professional, clear, and concise Release Management Document based on the following information:

### Release Notes:
{release_notes}

### Jira Tickets:
{linked_ticket_text}

Generate a professional summary of key changes, followed by detailed sections organized by category. Use Markdown format with headers (#), lists (-), and hyperlinks if needed. Do not include HTML tags.
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
    )
    response_text = response.choices[0].message.content.strip()
    response_text = re.sub(r"\(\s*\[\s*\]\)", "", response_text)    # elimina ([])
    response_text = re.sub(r"\(\s*\[\s*\]\s*\)", "", response_text) # elimina ([ ])
    response_text = re.sub(r"\(\s*\)", "", response_text)           # elimina ()
    response_text = re.sub(r"\s*\.\.+", ".", response_text)         # elimina puntos extra como ".."
    return response_text

def answer_question_with_gpt(question, release_info, history=None):
    if history is None:
        history = []

    version = release_info.get("version", "Unknown")
    body = release_info.get("body", "").strip()
    summary = release_info.get("summary", [])
    details = release_info.get("details", [])
    known_issues = release_info.get("known_issues", [])

    formatted_details = ""
    if details and isinstance(details[0], dict):
        formatted_details = "\n".join(f"- {d.get('ticket', '')}: {d.get('description', '')}" for d in details)
    elif isinstance(details, list):
        formatted_details = "\n".join(f"- {d}" for d in details)

    base_context = f"""
You are an AI release assistant. You help developers and QA understand software updates clearly.

Release: {version}

Full Release Notes (from GitHub or commit logs):
{body or "No notes available."}

Summary of Changes:
{', '.join(summary) if summary else 'None'}

Details:
{formatted_details if formatted_details else 'No detailed tickets provided.'}

Known Issues:
{', '.join(known_issues) if known_issues else 'None reported.'}
"""

    messages = [{"role": "system", "content": base_context}]
    messages.extend(history)
    messages.append({"role": "user", "content": question})

    response = client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        temperature=0.3,
    )

    return response.choices[0].message.content

def compare_releases_with_gpt(release_a: dict, release_b: dict) -> str:
    def format_details(details):
        if details and isinstance(details[0], dict):
            return "\n".join(f"- {d.get('ticket', '')}: {d.get('description', '')}" for d in details)
        elif isinstance(details, list):
            return "\n".join(f"- {d}" for d in details)
        else:
            return ""

    prompt = f"""
Compare the following two software releases.

Release A (version {release_a.get('version')}):
Summary:
{', '.join(release_a.get('summary', []))}

Detailed Notes:
{format_details(release_a.get('details', []))}

Known Issues:
{', '.join(release_a.get('known_issues', [])) or 'None'}

---

Release B (version {release_b.get('version')}):
Summary:
{', '.join(release_b.get('summary', []))}

Detailed Notes:
{format_details(release_b.get('details', []))}

Known Issues:
{', '.join(release_b.get('known_issues', [])) or 'None'}

Now produce a clear and structured comparison of the differences between these releases.
Focus on:
- New features or fixes added in B but not in A
- Issues resolved or introduced
- Key differences in ticket scope
- Any improvements or regressions
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return response.choices[0].message.content

def generate_professional_word(version_tag, content, path="release_ai_dashboard/static"):
    if not os.path.exists(path):
        os.makedirs(path)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    doc.add_heading(f"📘 Release Management Document - {version_tag}", level=1)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- [") and "](" in line:
            match = re.match(r"- \[(.*?)\]\((.*?)\): (.*)", line)
            if match:
                link_text, url, description = match.groups()
                p = doc.add_paragraph(style='List Bullet')
                add_hyperlink(p, f"{link_text}", url)
                p.add_run(f": {description}")
            else:
                doc.add_paragraph(line, style='List Bullet')
        elif line.startswith("- "):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)

    filename = f"release_{version_tag.replace('/', '-')}.docx"
    filepath = os.path.join(path, filename)
    doc.save(filepath)
    return filename

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
